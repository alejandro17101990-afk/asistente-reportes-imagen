import streamlit as st
from docx import Document
import speech_recognition as sr
from openai import OpenAI
import io

st.set_page_config(page_title="Beam AI Radiology", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&display=swap');
    .stApp { background-color: #0b0b0f; color: #e2e8f0; font-family: 'Inter', sans-serif; }
    [data-testid="stSidebar"] { background-color: #111116; border-right: 1px solid #1f1f2e; }
    .stButton>button { background: linear-gradient(135deg, #7c3aed 0%, #6d28d9 100%) !important; color: white !important; border-radius: 8px !important; border: 1px solid #8b5cf6 !important; padding: 0.6rem 1rem !important; font-weight: 500 !important; width: 100%; }
    div[data-testid="stVerticalBlock"] > div[style*="flex-direction: column;"] { background-color: #12121a; border: 1px solid #1f1f2e; border-radius: 12px; padding: 24px; }
    .stTextInput input, .stSelectbox div[data-baseweb="select"], .stTextArea textarea { background-color: #16161d !important; color: #f8fafc !important; border: 1px solid #2a2a35 !important; border-radius: 8px !important; }
    h1, h2, h3, h4 { color: #ffffff !important; font-weight: 600 !important; }
    hr { border-color: #1f1f2e !important; }
    </style>
    """, unsafe_allow_html=True)

def leer_word_con_tablas(file):
    doc = Document(file)
    contenido = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        contenido.append("\n[FORMATO DE TABLA REQUERIDO]")
        for row in table.rows:
            contenido.append(" | ".join([cell.text.replace('\n', ' ') for cell in row.cells]))
    return '\n'.join(contenido)

def generar_docx(texto_limpio):
    doc = Document()
    doc.add_paragraph(texto_limpio)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def transcribir_audio(audio_file):
    r = sr.Recognizer()
    with sr.AudioFile(audio_file) as source:
        try: return r.recognize_google(r.record(source), language="es-MX")
        except: return "[No se detectó voz clara]"

if 'reporte_generado' not in st.session_state: st.session_state.reporte_generado = ""

with st.sidebar:
    st.markdown("### ⚕️ Beam AI")
    try:
        api_key = st.secrets["deepseek_key"]
        st.success("Motor DeepSeek: Conectado 🟢")
    except:
        api_key = st.text_input("DeepSeek Key", type="password")
    
    st.divider()
    archivo_plantilla = st.file_uploader("Cargar Word (.docx)", type=["docx"])
    plantilla_contenido = leer_word_con_tablas(archivo_plantilla) if archivo_plantilla else ""
    reglas_usuario = st.text_area("Reglas de redacción:", height=150, placeholder="Ej. No usar gerundios. Agrupar la patología...")

if api_key:
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com", timeout=60.0)
    st.markdown("## ⚡ Workspace Radiológico")
    col_input, col_output = st.columns([1, 1.2], gap="large")

    with col_input:
        mod_col, sub_col = st.columns(2)
        with mod_col: modalidad = st.selectbox("Modalidad", ["Tomografía", "Resonancia", "Radiografía", "Ultrasonido", "PET-CT", "Procedimiento"])
        with sub_col: estilo = st.selectbox("Estilo", ["Conciso", "Académico", "Institucional"])
        
        audio_file = st.audio_input("Dictado Rápido")
        notas_texto = st.text_area("Hallazgos:", height=280, placeholder="Escribe aquí los hallazgos...")

        if st.button("✨ Generar Informe Clínico"):
            texto_dictado = transcribir_audio(audio_file) if audio_file else ""
            if not notas_texto and not texto_dictado:
                st.warning("⚠️ Ingresa hallazgos por texto o voz antes de generar.")
            else:
                prompt_sistema = f"Eres un Radiólogo de alta especialidad. Genera informe estructurado de {modalidad} en estilo {estilo}. Precisión anatómica absoluta. SI HAY TABLAS, LLÉNALAS. PLANTILLA: {plantilla_contenido if plantilla_contenido else 'Ninguna'}. REGLAS USUARIO: {reglas_usuario}"
                with st.spinner("🧠 Analizando hallazgos e integrando plantilla..."):
                    try:
                        response = client.chat.completions.create(
                            model="deepseek-chat",
                            messages=[{"role": "system", "content": prompt_sistema}, {"role": "user", "content": f"Notas: {notas_texto} | Dictado: {texto_dictado}"}],
                            temperature=0.2
                        )
                        st.session_state.reporte_generado = response.choices[0].message.content
                    except Exception as e: st.error(f"❌ Error: {str(e)}")

    with col_output:
        texto_final = st.text_area("Revisa y afina tu reporte:", value=st.session_state.reporte_generado, height=600)
        if st.session_state.reporte_generado:
            col_dl, col_ref = st.columns(2)
            with col_dl:
                st.download_button(label="📥 Descargar a Word", data=generar_docx(texto_final), file_name="Reporte_Radiologico.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col_ref:
                if st.button("🔄 Reformular Conclusión"):
                    with st.spinner("Reevaluando impresión diagnóstica..."):
                        try:
                            prompt_ref = f"Mejora ÚNICAMENTE la Impresión Diagnóstica agrupando diagnósticos. Mantén el resto igual. REPORTE:\n{texto_final}"
                            response_ref = client.chat.completions.create(
                                model="deepseek-chat", messages=[{"role": "system", "content": "Jefe de radiología corrigiendo."}, {"role": "user", "content": prompt_ref}], temperature=0.3)
                            st.session_state.reporte_generado = response_ref.choices[0].message.content
                            st.rerun()
                        except: st.error("Error al reformular.")
else: st.info("Configura tu credencial para inicializar el Workspace.")
